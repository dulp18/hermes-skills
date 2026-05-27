# Working Code Snippets

## Verify all files after filling

```python
import os
from docx import Document
from openpyxl import load_workbook

desktop = os.path.expanduser("~/Desktop")
files = ["file1.docx", "file2.xlsx"]

for fname in files:
    path = os.path.join(desktop, fname)
    size = os.path.getsize(path)
    if fname.endswith('.docx'):
        doc = Document(path)
        filled_cells = sum(1 for t in doc.tables for r in t.rows[1:] for c in r.cells if c.text.strip())
        print(f"[OK] {fname} — {size/1024:.1f}KB | {len(doc.tables)} tables | {filled_cells} filled cells")
    else:
        wb = load_workbook(path)
        ws = wb.active
        filled = sum(1 for row in ws.iter_rows(min_col=3) for c in row if c.value and str(c.value).strip())
        print(f"[OK] {fname} — {size/1024:.1f}KB | {ws.title} | {filled} filled rows")
```

## Batch fill Word tables from JSON

```python
import json
from docx import Document

with open('/tmp/data.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

doc = Document("path.docx")

# data key format: "T{table_index}_{row_index}"
# Example: data["T0_1"] = "text for table 0, row 1, cell 1"
for key, text in data.items():
    parts = key.split("_")
    table_idx = int(parts[0][1:])  # "T0" -> 0
    row_idx = int(parts[1])
    cell = doc.tables[table_idx].rows[row_idx].cells[1]
    for p in cell.paragraphs:
        p.clear()
    cell.paragraphs[0].text = text

doc.save("path.docx")
```

## Batch fill Excel from JSON

```python
import json
from openpyxl import load_workbook

with open('/tmp/data.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

wb = load_workbook("path.xlsx")
ws = wb.active

# data keys are row numbers as strings
for row_str, text in data.items():
    row = int(row_str)
    ws.cell(row=row, column=3).value = text

wb.save("path.xlsx")
```
